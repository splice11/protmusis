#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate printable A4 team-name cards for a trivia night.

Produces two .docx files in ./team_cards/:
  1. table_signs.docx  - A4 landscape, one large team name per page (table tents).
  2. draw_slips.docx    - A4 portrait, grid of small slips (5 per team) to draw.
"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# A4 dimensions
A4_SHORT = Cm(21.0)
A4_LONG = Cm(29.7)

FONT_NAME = "Calibri"

# 10 team names (Lithuanian rivers/lakes) in the required order.
TEAMS = [
    "Asveja",
    "Tauragnas",
    "Dusia",
    "Sartai",
    "Plateliai",
    "Neris",
    "Nemunas",
    "Venta",
    "Merkys",
    "Šešupė",
]

# Per-team font size (pt) for the big table signs. Longer names get smaller
# sizes so they still fit on one line within ~1cm landscape margins.
SIGN_FONT_SIZES = {
    "Asveja": 150,
    "Tauragnas": 115,
    "Dusia": 150,
    "Sartai": 150,
    "Plateliai": 120,
    "Neris": 150,
    "Nemunas": 120,
    "Venta": 150,
    "Merkys": 140,
    "Šešupė": 150,
}

SLIPS_PER_TEAM = 5  # 5 members per team


def _set_run_font(run, name=FONT_NAME):
    """Force the font name across all script ranges (incl. East Asian/cs)."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _set_cell_borders(cell, sz=12, val="dashed", color="888888"):
    """Apply visible cut borders to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)


def _vertical_center_cell(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_row_border(row, edge, sz=12, val="dashed", color="888888"):
    """Apply a single visible border to every cell on one edge of a row."""
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def _set_row_exact_height(row, height):
    row.height = height
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def build_table_signs(path):
    """A4 landscape table tents, one team per page (10 pages).

    Each page is split in half by a horizontal fold line. The team name sits
    in the *bottom* half only; the top half is blank. Fold along the dashed
    line into an A-frame and the name shows upright on one face.
    """
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = A4_LONG
    section.page_height = A4_SHORT
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(1.0))

    # Usable area: 29.7 - 2 = 27.7 cm wide, 21 - 2 = 19 cm tall. Keep the two
    # halves just under half each so rounding never spills a blank page.
    half_height = Cm(9.3)
    full_width = Cm(27.5)

    for idx, team in enumerate(TEAMS):
        # Two-row, single-column table filling the page: blank top half, then
        # the name centred in the bottom half. A dashed border between the
        # rows marks the fold line.
        table = doc.add_table(rows=2, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        top, bottom = table.rows[0], table.rows[1]
        _set_row_exact_height(top, half_height)
        _set_row_exact_height(bottom, half_height)

        for row in (top, bottom):
            row.cells[0].width = full_width

        # Fold line across the middle of the page.
        _set_row_border(top, "bottom")

        cell = bottom.cells[0]
        cell.width = full_width
        _vertical_center_cell(cell)
        cell_p = cell.paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_p.add_run(team)
        run.bold = True
        run.font.size = Pt(SIGN_FONT_SIZES[team])
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        _set_run_font(run)

        if idx != len(TEAMS) - 1:
            # Page break between teams; no trailing blank page.
            doc.add_page_break()

    doc.save(path)


def build_draw_slips(path):
    """A4 portrait grid of small slips, 5 per team (50 total)."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = A4_SHORT
    section.page_height = A4_LONG
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(1.0))

    # Build the full list of slips: 5 of each team. Interleave so each page
    # has variety, though ordering is unimportant.
    slips = []
    for i in range(SLIPS_PER_TEAM):
        for team in TEAMS:
            slips.append(team)
    # slips now has 50 entries, exactly 5 of each team.

    # Small, uniform slips: a dense grid of equal cells. 5 columns x 10 rows
    # fits all 50 slips (5 of each of the 10 teams) on a single A4 page.
    COLS = 5
    ROWS_PER_PAGE = 10
    PER_PAGE = COLS * ROWS_PER_PAGE  # 50
    ROW_HEIGHT = Cm(2.5)
    CELL_WIDTH = Cm(3.6)

    total_pages = (len(slips) + PER_PAGE - 1) // PER_PAGE

    for page in range(total_pages):
        page_slips = slips[page * PER_PAGE:(page + 1) * PER_PAGE]
        rows = (len(page_slips) + COLS - 1) // COLS

        table = doc.add_table(rows=rows, cols=COLS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for r in range(rows):
            _set_row_exact_height(table.rows[r], ROW_HEIGHT)

        for i, team in enumerate(page_slips):
            r, c = divmod(i, COLS)
            cell = table.cell(r, c)
            cell.width = CELL_WIDTH
            _set_cell_borders(cell)
            _vertical_center_cell(cell)
            cell_p = cell.paragraphs[0]
            cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell_p.add_run(team)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            _set_run_font(run)

        if page != total_pages - 1:
            doc.add_page_break()

    doc.save(path)


def main():
    out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "team_cards")
    os.makedirs(out_dir, exist_ok=True)

    signs_path = os.path.join(out_dir, "table_signs.docx")
    slips_path = os.path.join(out_dir, "draw_slips.docx")

    build_table_signs(signs_path)
    build_draw_slips(slips_path)

    print(f"Wrote: {signs_path}")
    print(f"Wrote: {slips_path}")


if __name__ == "__main__":
    main()
