#!/usr/bin/env python3
"""Generate a clean, printable A4 answer key for the Trivia Night.

One compact table per round — question number, a brief version of the
question, the correct answer and the points — so a marker can quickly
correct what each team wrote. Black on white, fits on a couple of A4 pages.

Usage:  pip install python-docx && python3 generate_answer_sheet.py
Output: Trivia_Night_2026-06-16_answer_key.docx
"""

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x1E, 0x3C, 0x86)       # EU blue
HEAD_BG = "1E3C86"
ALT_BG = "EEF1F7"

# round title -> list of (q_no, brief question, answer, points)
ROUNDS = [
    ("Round 1 — Europe & Fun Facts", [
        ("1", "Which European island changes country every six months?",
         "B — Pheasant Island", "1"),
        ("2", "Which EU symbol was unveiled decades ago (recently ~70)?",
         "B — The European flag", "1"),
        ("3", "Which major international award did the EU win in 2012?",
         "B — The Nobel Peace Prize", "1"),
        ("4", "Which euro banknote was nicknamed “Bin Laden”?",
         "D — The €500 note", "1"),
        ("5", "Which EU country reached the World Cup first time since 1998?",
         "D — Austria", "1"),
        ("6", "Umaru Dikko, 1984 — what went wrong with the diplomatic crate?",
         "A — The crate was improperly labelled as diplomatic baggage", "1"),
        ("7", "What made the Latvia-born inventor's blue jeans famous?",
         "B — Metal rivets", "1"),
        ("8", "Which part of the Artemis II spacecraft did ESA build?",
         "C — The life-support systems (European Service Module)", "1"),
    ]),
    ("Round 2 — Lithuania", [
        ("1", "2006 Eurovision — the song that declared victory in advance?",
         "“We Are The Winners” (LT United)", "1"),
        ("2", "What did the two Lithuanian mountaineers scatter from Everest?",
         "Amber", "1"),
        ("3", "Which country out-produced 1938 Lithuania in butter & bacon?",
         "Denmark", "1"),
        ("4", "Which mythical flower is sought on Midsummer night?",
         "The fern flower", "1"),
        ("5", "What did Lithuanian women gain in 1918?",
         "Voting rights / the right to vote", "1"),
        ("6", "Which souvenirs were thrown overboard during the storm?",
         "Egyptian mummies", "1"),
        ("7", "Which approaching event did the cepelinai promo reflect?",
         "Joining the European Union (2004)", "1"),
        ("8", "What colour is the soup of the beloved Vilnius festival?",
         "Pink (šaltibarščiai)", "1"),
    ]),
    ("Round 3 — Environment & Nature", [
        ("1", "120 of what did Metallica's Antarctic audience use?",
         "Headphones", "1"),
        ("2", "Which animals did the three “Trimates” study?",
         "Great apes", "1"),
        ("3", "Which EU law hit the news after soldiers died in a bog?",
         "The Nature Restoration Law", "1"),
        ("4", "Which “ecosystem engineer” builds wetlands?",
         "The beaver", "1"),
        ("5", "Which EU state is ~three-quarters covered by forest?",
         "Finland", "1"),
        ("6", "Which river runs through/borders the most European countries?",
         "The Danube", "1"),
    ]),
    ("Round 4 — Music (Name the Connection)", [
        ("1", "Which device, dreamed of since antiquity, was hidden by OOO?",
         "A submarine  (clue: “Yellow Submarine”)", "1"),
        ("2", "EU initiative for world-leading ocean intelligence by 2035?",
         "“OceanEye”  (clue: “Ocean Eyes”)", "1"),
        ("3", "Which 2016 film, renamed across Europe, is about the sea?",
         "“Moana” / “Vaiana”  (clue: “How Far I'll Go”)", "1"),
        ("4", "Who made a 1997 blockbuster to visit the real wreck?",
         "James Cameron  (clue: “My Heart Will Go On”)", "1"),
        ("5", "The living, the dead, and…? (round's hidden theme)",
         "…and those at sea  (clue: Čiurlionis, “Jūra”)", "1"),
    ]),
    ("Bonus — The Brussels Bubble", [
        ("1", "On which file did Romania praise Lithuania for being so active?",
         "The Waste Framework Directive", "1"),
        ("2", "What did Lithuania's DPR say at the split Coreper I, and file?",
         "“Scrutiny reservation” — Nature Restoration Law", "2"),
        ("3", "How did the USSR try to derail the Baltic Star cruise?",
         "A (false) rumour that a bomb was on board", "1"),
    ]),
    ("Tie-breaker (only if level)", [
        ("1", "Who so admired tiny St Anne's Church in 1812 Vilnius?",
         "Napoleon Bonaparte", "—"),
    ]),
]


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def _cell(cell, text, *, bold=False, color=INK, size=10.5,
          align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = 1  # center
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color


doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.PORTRAIT
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(1.5)
sec.left_margin = sec.right_margin = Cm(1.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Trivia Night — Answer Key")
r.font.name = "Calibri"
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
r = sub.add_run("16 June 2026 · Permanent Representation of Lithuania to the EU"
                "   ·   for marking team answers")
r.font.size = Pt(10)
r.font.italic = True
r.font.color.rgb = MUTED
sub.paragraph_format.space_after = Pt(6)

WIDTHS = [Cm(0.9), Cm(8.6), Cm(7.4), Cm(1.0)]

for rtitle, rows in ROUNDS:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(3)
    r = h.add_run(rtitle)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for c, txt, al in zip(
            hdr, ["#", "Question", "Answer", "Pts"],
            [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
             WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]):
        _shade(c, HEAD_BG)
        _cell(c, txt, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10.5,
              align=al)
    for i, (qno, q, a, pts) in enumerate(rows):
        cells = table.add_row().cells
        if i % 2 == 1:
            for c in cells:
                _shade(c, ALT_BG)
        _cell(cells[0], qno, align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
        _cell(cells[1], q, italic=True, color=INK)
        _cell(cells[2], a, bold=True, color=INK)
        _cell(cells[3], pts, align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    _set_widths(table, WIDTHS)

OUT = "Trivia_Night_2026-06-16_answer_key.docx"
doc.save(OUT)
print(f"Saved {OUT}")
