#!/usr/bin/env python3
"""Generate the run-of-show Word document for the Trivia Night deck.

The document mirrors the final order of Trivia_Night_2026-06-16.pptx —
rounds, questions, options, answers, did-you-know facts and host notes —
and closes with a summary of what changed relative to the original draft
(`Klausimai protmūšiui (2).docx`).

Usage:  pip install python-docx && python3 generate_rundown.py
Output: Trivia_Night_2026-06-16_run_of_show.docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0xA8, 0x6A, 0x14)
MUTED = RGBColor(0x6E, 0x75, 0x82)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def para(text="", bold=False, italic=False, color=None, style_name=None,
         space_after=6):
    p = doc.add_paragraph(style=style_name)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def labelled(label, body, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(label + "  ")
    r.bold = True
    if color:
        r.font.color.rgb = color
    r2 = p.add_run(body)
    r2.italic = italic
    if color:
        r2.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def question(no, total, q, options=None, correct=None, answer=None,
             fact=None, hint=None, video=None, notes=None, points=None):
    head = f"Question {no} of {total}" + (f" · {points}" if points else "")
    doc.add_heading(head, level=3)
    para(q)
    if options:
        for letter, opt in options:
            p = doc.add_paragraph()
            r = p.add_run(f"{letter}.  {opt}")
            if letter == correct:
                r.bold = True
                p.add_run("   ← correct").italic = True
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(18)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
    if hint:
        labelled("Hint:", hint, italic=True)
    if video:
        labelled("Video:", video)
    labelled("Answer:", answer, color=ACCENT)
    if fact:
        labelled("Did you know?", fact)
    if notes:
        labelled("Host notes:", notes, italic=True, color=MUTED)


# ------------------------------------------------------------------ cover --
title = doc.add_heading("Trivia Night — run of show", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
para("Brussels · 16 June 2026 · Permanent Representation of Lithuania to "
     "the European Union · Environment Team", bold=True)
para("This document follows the final slide deck "
     "(Trivia_Night_2026-06-16.pptx) in presentation order. It was built "
     "from the question draft in “Klausimai protmūšiui (2).docx”; the changes "
     "made relative to the draft are summarised at the end.", italic=True,
     color=MUTED)

doc.add_heading("Opening — how tonight works", level=1)
para("Team draw at the door: every guest draws a slip with the name of a "
     "Lithuanian river or lake — that is their team.")
para("Four rounds plus a bonus round. One point per question unless "
     "marked otherwise.")
para("Teams write their answers down; answers are revealed and scored "
     "after each round.")
para("Fair play: no phones — diplomatic immunity does not cover googling.")

# ----------------------------------------------------------------- round 1 --
doc.add_heading("Round 1 — Europe & Fun Facts", level=1)
para("Multiple choice · 8 questions · 1 point each", italic=True,
     color=MUTED)

question(
    1, 8, "Which European island changes sovereignty every six months?",
    options=[("A", "Heligoland"), ("B", "Pheasant Island"), ("C", "Jersey"),
             ("D", "Bornholm")],
    correct="B", answer="B — Pheasant Island",
    fact="Under the 1659 Treaty of the Pyrenees, France and Spain share "
         "this tiny uninhabited island in the Bidasoa river. Administration "
         "alternates between the two countries every six months — the "
         "world's oldest condominium.")

question(
    2, 8, "Which EU symbol was publicly unveiled about 40 years ago?",
    options=[("A", "Euro coins"), ("B", "European flag"),
             ("C", "European anthem"), ("D", "Schengen passport")],
    correct="B", answer="B — The European flag",
    fact="In May 1986 the European flag was raised for the first time "
         "outside the Berlaymont building — twelve gold stars in a circle, "
         "a symbol of unity, unchanged ever since.")

question(
    3, 8, "What major international award did the European Union receive "
          "in 2012?",
    options=[("A", "Sakharov Prize"), ("B", "Nobel Peace Prize"),
             ("C", "Charlemagne Prize"), ("D", "Right Livelihood Award")],
    correct="B", answer="B — The Nobel Peace Prize",
    fact="Awarded to the EU in 2012 for over six decades of advancing "
         "peace, reconciliation, democracy and human rights in Europe.")

question(
    4, 8, "Which euro banknote was nicknamed “Bin Laden”?",
    options=[("A", "€50"), ("B", "€100"), ("C", "€200"), ("D", "€500")],
    correct="D", answer="D — The €500 note",
    fact="The ECB stopped issuing the €500 note in 2019 over its "
         "popularity with money launderers — and because, like its "
         "namesake, everyone knew what it looked like but almost no one "
         "had ever seen one.")

question(
    5, 8, "Which EU country is playing in this year's football World Cup "
          "for the first time since 1998?",
    options=[("A", "Belgium"), ("B", "Czechia"), ("C", "Sweden"),
             ("D", "Austria")],
    correct="D", answer="D — Austria",
    fact="Austria are back at the World Cup finals for the first time since "
         "France 1998 — a long wait finally over for the side in red and "
         "white.")

question(
    6, 8, "In 1984, former Nigerian minister Umaru Dikko was kidnapped in "
          "London. His captors planned to smuggle him out of the UK in a "
          "diplomatic crate, relying on the Vienna Convention rule that "
          "diplomatic bags cannot be opened or detained. What went wrong?",
    options=[("A", "The crate was improperly labelled as diplomatic "
                   "baggage"),
             ("B", "Dikko did not fit inside the crate"),
             ("C", "The aircraft meant to transport him never arrived"),
             ("D", "Customs officers had never heard of the Vienna "
                   "Convention")],
    correct="A", answer="A — The crate was improperly labelled",
    fact="With no official diplomatic markings, customs officers at "
         "Stansted were entitled to open the crate — and found Dikko "
         "unconscious, accompanied by an anaesthetist.",
    video="https://www.youtube.com/watch?v=N83Idy9IOmU (the Umaru Dikko "
          "story — shown with the answer)")

question(
    7, 8, "One of the inventors of modern blue jeans was born in "
          "present-day Latvia. He developed a way to make work trousers "
          "much more durable and partnered with Levi Strauss to patent the "
          "idea. What simple innovation made the trousers famous?",
    options=[("A", "A zipper"), ("B", "Metal rivets"),
             ("C", "Waterproof fabric"), ("D", "Belt loops")],
    correct="B", answer="B — Metal rivets",
    fact="Jacob Davis, a tailor born in Riga in 1831, reinforced the "
         "stress points of work trousers with copper rivets. He and Levi "
         "Strauss patented the idea in 1873 — and blue jeans were born.")

question(
    8, 8, "Earlier this year the world watched the Artemis II crew loop "
          "around the Moon — the first crewed lunar voyage in over 50 "
          "years. Which part of their spacecraft was built by the European "
          "Space Agency?",
    options=[("A", "The rocket boosters"), ("B", "The crew quarters"),
             ("C", "The life-support systems"), ("D", "The heat shield")],
    correct="C", answer="C — The life-support systems",
    fact="ESA built the European Service Module that powers NASA's Orion "
         "spacecraft — providing propulsion, electricity, and the air, "
         "water and temperature control that keep the crew alive.")

# ----------------------------------------------------------------- round 2 --
doc.add_heading("Round 2 — Lithuania", level=1)
para("8 questions · 1 point each", italic=True, color=MUTED)

question(
    1, 8, "Every year Lithuania plans to win Eurovision — although "
          "officially we never have. In 2006, however, Lithuania tried a "
          "different strategy: entering a song that simply declared "
          "victory before the contest was even over. What was the title "
          "of the song?",
    answer="“We Are The Winners” (LT United)",
    fact="Honourable mention: Lithuanian-German singer Lena Valaitis took "
         "second place at Eurovision 1981 — proof that we keep looking "
         "for that victory by all possible means.",
    video="https://www.youtube.com/watch?v=DBAdOlQPbwg",
    notes="The clip plays on the answer slide (the question keeps just the "
          "Eurovision logo, so it isn't a give-away). Decide which part of "
          "the clip to show. Mention Lena Valaitis as a fun fact, not a "
          "question.")

question(
    2, 8, "Two Lithuanian mountaineers carried a symbolic national object "
          "to the summit of Mount Everest and scattered it from the "
          "world's highest peak. What was it?",
    answer="Amber",
    fact="Vladas Vitkauskas (1995) and Saulius Vilius (2003) both carried "
         "Baltic amber to the top of the world and scattered it from the "
         "summit.",
    notes="Tell the full story and the exact years out loud; the slide "
          "stays short.")

question(
    3, 8, "When Lithuanians learn about the interwar period, one "
          "statistic is proudly repeated in schools: in 1938, Lithuania "
          "ranked second in Europe in butter and bacon production — "
          "despite being one of the poorest countries on the continent. "
          "Which country ranked first?",
    hint="Think of a recent Council Presidency.",
    answer="Denmark",
    fact="A statistic still celebrated in Lithuanian classrooms. Denmark, "
         "for its part, no longer wants to be the EU's bacon factory — as "
         "Politico once put it.")

question(
    4, 8, "Across Europe people celebrate Midsummer with bonfires. In "
          "Lithuania, one tradition involves searching for a mythical "
          "flower that supposedly blooms only on that night. Which "
          "flower?",
    answer="The fern flower",
    fact="According to tradition, the fern blooms only on Midsummer "
         "night — whoever finds it gains happiness and wisdom. Botanists "
         "remain unconvinced.")

question(
    5, 8, "1918 was a big year for Lithuania — also for democracy. "
          "Lithuanian women gained something that many women in Western "
          "Europe had to wait decades for. What was it?",
    answer="Voting rights",
    fact="Lithuanian women gained the vote on 2 November 1918, when the "
         "provisional constitution enshrined equal suffrage — ahead of "
         "much of Western Europe. A commemorative stamp marked the "
         "centenary.",
    notes="Alternative version from the draft: show the centenary stamp "
          "(details removed) and ask what occasion it marks.")

question(
    6, 8, "A 16th-century Lithuanian nobleman travelled through Egypt and "
          "bought several unusual souvenirs. During a storm at sea, "
          "terrified sailors blamed the bad weather on them and threw "
          "them overboard. What were the souvenirs?",
    answer="Egyptian mummies",
    fact="Mikalojus Kristupas Radvila the Orphan brought two mummies back "
         "from his pilgrimage. When storms battered the ship, the crew "
         "blamed the cargo — and overboard they went. His travel diary "
         "became a European bestseller.")

question(
    7, 8, "In a humorous Lithuanian promotional video from the early "
          "2000s, traditional dishes such as cepelinai are discussed with "
          "the idea that they might one day become popular across Europe "
          "as street food. The video reflects a moment when Lithuania was "
          "preparing for a major change. What event was approaching?",
    answer="Joining the European Union",
    fact="Lithuania joined the EU on 1 May 2004, in the largest "
         "enlargement in the Union's history — ten countries at once. The "
         "cepelinai street-food revolution is still pending.",
    video="https://www.youtube.com/watch?v=YgvHcenDYcU",
    notes="Add English subtitles to the clip. Consider whether to keep "
          "“early 2000s” in the wording.")

question(
    8, 8, "A few weeks ago a beloved Vilnius festival was held for the 4th "
          "time. Its organisers wrote: “The most beautiful love story in "
          "the world? No — we mean the Lithuanian love for a cold OOO "
          "soup… it's not just a soup, it's a way of life.” What colour is "
          "this soup?",
    hint="In the year of the very first festival, this colour was a "
         "worldwide craze.",
    answer="Pink",
    fact="Šaltibarščiai — the electric-pink cold beet soup — is Lithuania's "
         "unofficial summer flag. The Vilnius “Pink Soup Fest” has become a "
         "hit: last year more than 90,000 fans turned the city pink.")

# ----------------------------------------------------------------- round 3 --
doc.add_heading("Round 3 — Environment & Nature", level=1)
para("6 questions · 1 point each", italic=True, color=MUTED)

question(
    1, 6, "In 2013 Metallica became the first band to perform on all "
          "seven continents. During their Antarctic concert, the audience "
          "used 120 of what?",
    answer="Headphones",
    fact="To comply with Antarctic environmental rules, the band played "
         "without amplifiers — the audience of about 120 listened through "
         "headphones. The concert was fittingly called “Freeze 'Em All”.")

question(
    2, 6, "Three famous scientists — Birutė Galdikas, Dian Fossey and "
          "Jane Goodall — devoted their careers to studying these animals "
          "in the wild. Collectively, they became known as the "
          "“Trimates”. What group of animals are these?",
    answer="Great apes",
    fact="Goodall (chimpanzees), Fossey (gorillas) and Galdikas "
         "(orangutans) were recruited by Louis Leakey — hence also "
         "“Leakey's Angels”. Galdikas, of Lithuanian descent, still works "
         "in Borneo. “Orangutan” means “person of the forest”.")

question(
    3, 6, "Which EU law unexpectedly entered the headlines after the "
          "deaths of four US soldiers, whose vehicle sank in a military "
          "training area near the Lithuanian–Belarusian border?",
    answer="The Nature Restoration Law",
    fact="Commentary on the 2025 tragedy near Pabradė linked the swampy "
         "terrain to wetlands restored under EU environmental policy — "
         "putting the Nature Restoration Law unexpectedly in the news.")

question(
    4, 6, "This animal is often called an “ecosystem engineer” because it "
          "creates wetlands that benefit countless other species. What "
          "animal is it?",
    answer="The beaver",
    fact="Beaver dams create wetlands that store water, filter pollution "
         "and host countless species. Lithuania's beaver population has "
         "grown from near extinction to one of the densest in Europe.")

question(
    5, 6, "We all love forests — yet one EU Member State is so densely "
          "wooded that roughly three-quarters of its land area is covered "
          "by forest, the highest share in the whole Union. Which Member "
          "State is it?",
    answer="Finland",
    fact="About three-quarters of Finland's land area is forest — the "
         "highest share of any EU country, and the backdrop to its image "
         "as a land of endless woods and a thousand lakes.")

question(
    6, 6, "We talk a lot about water resilience. Which river flows "
          "through — or forms the border of — the greatest number of "
          "countries in Europe?",
    answer="The Danube",
    fact="Europe's most international river runs about 2,850 km from "
         "Germany's Black Forest to the Black Sea, flowing through or "
         "bordering 10 countries — more than any other river in Europe.")

# ----------------------------------------------------------------- round 4 --
doc.add_heading("Round 4 — Music (Name the Connection)", level=1)
para("5 riddles · 1 point each · click the embedded song excerpt on each "
     "slide to hear the clue", italic=True, color=MUTED)
para("Hidden theme of the round: every answer is tied to the sea — a nod "
     "to the EU's upcoming Ocean Pact (and to Čiurlionis's “Jūra”). Reveal "
     "the link at the end for an optional extra point.", italic=True,
     color=MUTED)

question(
    1, 5, "Drawings and diagrams dating back centuries show that using OOO "
          "has always been a human dream. Records from 415 BC, at the Siege "
          "of Syracuse, even describe its organised military use. What "
          "device hides behind OOO?",
    answer="A submarine",
    fact="Clue song: The Beatles' “Yellow Submarine” (1966). Humans have "
         "dreamed of travelling underwater since antiquity — and every "
         "answer in this round lives at sea.")

question(
    2, 5, "The European Commission recently launched a new initiative that, "
          "by 2035, should make the EU the world's leading provider of OOO "
          "intelligence. What is the initiative called?",
    answer="“OceanEye”",
    fact="Clue song: Billie Eilish's “Ocean Eyes” (2015). The EU's push for "
         "world-leading ocean intelligence keeps the round firmly at sea.")

question(
    3, 5, "This 2016 animated movie centres on saving the dying heart of "
          "this round's theme. Its 2024 sequel, like the original, had to "
          "be renamed in some European countries. What is the film's title?",
    answer="“Moana” (a.k.a. “Vaiana”)",
    fact="Clue song: “How Far I'll Go” from Disney's “Moana” (2016), "
         "renamed “Vaiana” across much of Europe — and all about the call "
         "of the ocean.")

question(
    4, 5, "It is said the main reason he made his 1997 blockbuster was so "
          "that he could visit the real thing on the ocean floor. Who is "
          "he?",
    answer="James Cameron",
    fact="Clue song: “My Heart Will Go On” from “Titanic” (1997). Director "
         "James Cameron, an obsessive deep-sea explorer, has dived to the "
         "wreck many times.")

question(
    5, 5, "The Ancient Greeks said there are three kinds of people: the "
          "living, the dead, and OOO. If you've worked out the theme of "
          "this round, the answer shouldn't be lost on you.",
    answer="“…and those at sea”",
    fact="Clue song: Čiurlionis's “Jūra” (“The Sea”). The old saying — the "
         "living, the dead, and those at sea — and the thread through every "
         "answer tonight: the sea, and the EU's upcoming Ocean Pact.")

# ------------------------------------------------------------- bonus round --
doc.add_heading("Bonus Round — The Brussels Bubble", level=1)
para("Insider questions · for those who were in the room", italic=True,
     color=MUTED)

question(
    1, 3, "During negotiations on which file did Violeta Dragu (Romania) "
          "say: “I have never seen Lithuania being so active during the "
          "negotiations”?",
    answer="The Waste Framework Directive",
    notes="Can be made multiple choice to make it easier — or harder, by "
          "also asking what the problem was.")

question(
    2, 3, "During a Coreper I meeting, the Council was split on a "
          "legislative file, with both sides carefully counting votes. "
          "The Deputy Permanent Representative of Lithuania intervened "
          "and said… what? And which file was being discussed?",
    points="2 points",
    answer="“Scrutiny reservation” — on the Nature Restoration Law",
    fact="A scrutiny reservation lets a member state hold its position "
         "while procedures are completed back home — a small phrase that "
         "can pause a finely balanced vote.",
    notes="Optional hint: give the exact date — the 14 June Coreper I (or "
          "the earlier one).")

question(
    3, 3, "1985: Lithuanian, Latvian and Estonian dissidents organised "
          "the Peace and Freedom Cruise aboard the ship Baltic Star — "
          "sailing from Stockholm along the edge of Soviet territorial "
          "waters and ending with a widely reported demonstration in "
          "Helsinki. The Soviet Union tried to derail the voyage. What "
          "did it do?",
    answer="It spread a rumour that a bomb was on board",
    fact="The rumour failed: the cruise went ahead, and the Helsinki "
         "demonstration was reported across the Scandinavian and wider "
         "European press — a loud reminder that the Baltic states had "
         "not been forgotten.",
    notes="Background: per Tomas Venclova, Aleksandras Štromas was the "
          "cruise's spiritus movens; the Baltic Tribunal also took place "
          "in 1985. Keep names off the slide.")

# ----------------------------------------------------------- tie-breaker --
doc.add_heading("Tie-breaker — if we need it", level=1)
para("One question to settle a draw · used only if teams are level after "
     "the bonus round", italic=True, color=MUTED)

question(
    1, 1, "When his army passed through Vilnius in 1812, legend says he so "
          "admired the tiny Church of St Anne that he wished he could carry "
          "it back to Paris in the palm of his hand. Near Kaunas, a hill is "
          "still called his “hat”. Who was he?",
    points="tie-breaker",
    answer="Napoleon Bonaparte",
    fact="As the Grande Armée marched east in 1812, Napoleon is said to "
         "have wished he could carry St Anne's Church back to Paris in his "
         "palm. A hill near Kaunas is still known as “Napoleon's hat”.")

doc.add_heading("Closing", level=1)
para("Thank you for playing — count the points and announce the winning "
     "team.")

# ------------------------------------------------- changes vs. the draft --
doc.add_page_break()
doc.add_heading("Changes relative to “Klausimai protmūšiui (2).docx”", level=1)

para("Structure and format", bold=True)
for line in [
    "One slide per question, followed by one slide with the answer — as "
    "proposed in the brief; everything translated into English for an "
    "international audience.",
    "Questions regrouped into balanced themed rounds: Round 1 — Europe & "
    "Fun Facts (8 multiple-choice), Round 2 — Lithuania (8), Round 3 — "
    "Environment & Nature (6), Round 4 — Music / Name the Connection (5 "
    "riddles with song clues), plus a Brussels Bubble bonus round (3 "
    "insider questions; the Coreper question is worth 2 points) and a "
    "one-question tie-breaker (Napoleon / St Anne's Church).",
    "Team assignment formalised: slips with Lithuanian rivers and lakes "
    "drawn at the door; a rules slide opens the evening.",
    "Production notes, source links and alternative question versions "
    "from the draft live in the speaker notes of each slide.",
]:
    doc.add_paragraph(line, style="List Bullet")

para("Content decisions", bold=True)
for line in [
    "Round 1 expanded to eight multiple-choice questions, adding the World "
    "Cup / Austria question and the Artemis II / ESA service-module "
    "question from the latest draft.",
    "Round 3 expanded to six questions, adding the most-forested Member "
    "State (Finland) and the most-international river (the Danube).",
    "Round 2's pink question reworked from a multiple-choice “which dish” "
    "into the open “Vilnius Pink Soup Fest” question (answer: pink).",
    "Round 4 reframed from “name this song” into five riddles whose "
    "answers all relate to the sea; each slide has a matching embedded "
    "song excerpt to play as the clue.",
    "A one-question tie-breaker added (Napoleon and St Anne's Church) for "
    "use only if teams finish level.",
    "Audio and video embedded to play inline in PowerPoint (no YouTube "
    "links): the five music-round excerpts, plus LT "
    "United “We Are The Winners” (excerpt to be chosen), the EU-accession "
    "promo (English subtitles still needed) and the Umaru Dikko story on "
    "its answer slide. Original YouTube URLs kept in the speaker notes as "
    "a backup.",
    "Lena Valaitis (2nd place, Eurovision 1981) kept as a fun fact on the "
    "answer slide rather than a separate question.",
    "Real photographs collected for nearly every question and answer "
    "(see photos/CREDITS.md for attribution); icons remain only where no "
    "suitable photo exists.",
    "Answer-giveaway icons removed from the Round 3 divider (beaver and "
    "headphones replaced with deer and tree).",
]:
    doc.add_paragraph(line, style="List Bullet")

para("Draft questions parked (not on slides — to discuss)", bold=True)
for line in [
    "CBAM question about nails (needs wording).",
    "The Schumann Show / the institutions (video).",
    "One question per member state / quick-fire round idea.",
    "Questions voiced by the Minister and Selemonas Paltanavičius "
    "(incl. the Belém fire).",
    "Vytis; the Pacai–Pazzi family link; Bona Sforza.",
    "German street in Vilnius.",
    "Kaunas modernism / Vilnius baroque.",
    "Adamkus 100 / Attenborough 100; the bears offered by Slovenia.",
    "Baltic Way symbols in the European Parliament.",
]:
    doc.add_paragraph(line, style="List Bullet")

OUT = "Trivia_Night_2026-06-16_run_of_show.docx"
doc.save(OUT)
print(f"Saved {OUT}")
